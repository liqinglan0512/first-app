"""Build the public v1.3.2 DOCX manual from a reviewable Markdown source.

The generated DOCX intentionally leaves TOC, cross-reference, caption-number,
page-number, and equation build-up work to Microsoft Word. Run
``tools/export_public_manual.ps1`` after this builder to materialize those
fields and export the matching public PDF.
"""

from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "manual" / "v1.3.2" / "manual-source.md"
DEFAULT_DOCX = (
    ROOT
    / "docs"
    / "manual"
    / "v1.3.2"
    / "Computational-Mechanics-Solver-v1.3.2-用户与技术说明书.docx"
)
BACKGROUND = ROOT / "web" / "welcome-bg.jpg"
COVER_IMAGE = ROOT / "docs" / "manual" / "v1.3.2" / "assets" / "cover-v1.3.2.png"

COVER_WIDTH_PX = 2480
COVER_HEIGHT_PX = 3508
SIMHEI_FONT = Path("C:/Windows/Fonts/simhei.ttf")
TIMES_BOLD_FONT = Path("C:/Windows/Fonts/timesbd.ttf")

A4_WIDTH = Mm(210)
A4_HEIGHT = Mm(297)
BODY_WIDTH = Mm(160)

CN_BODY = "宋体"
CN_HEADING = "黑体"
LATIN_BODY = "Times New Roman"
MATH_FONT = "Cambria Math"

INK = "1C2730"
MUTED = "5E6B73"
ACCENT = "176B77"
ACCENT_LIGHT = "E5F0F1"
LINE = "AEBFC4"
WARNING = "8A3E2F"
WARNING_FILL = "F8ECE8"

REF_PATTERN = re.compile(r"\{\{(eq|fig|table):([A-Za-z0-9_-]+)\}\}")
INLINE_PATTERN = re.compile(
    r"(\{\{(?:eq|fig|table):[A-Za-z0-9_-]+\}\}|\*\*[^*]+\*\*|`[^`]+`)"
)
IMAGE_PATTERN = re.compile(
    r"^!\[(?P<caption>.+?)\]\((?P<path>.+?)\)"
    r"(?:\{#(?P<id>[A-Za-z0-9_-]+)(?:\s+width=(?P<width>[0-9.]+)mm)?\})?$"
)
TABLE_DIRECTIVE = re.compile(
    r"^:::table\s+id=(?P<id>[A-Za-z0-9_-]+)\s+title=(?P<title>.+)$"
)
EQUATION_DIRECTIVE = re.compile(
    r"^:::equation\s+id=(?P<id>[A-Za-z0-9_-]+)\s+text=(?P<text>.+)$"
)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 100, start: int = 110, bottom: int = 100, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_table_borders(table, color: str = LINE, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _set_table_geometry(table, widths_twips: list[int]) -> None:
    total = sum(widths_twips)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "110")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_twips[min(index, len(widths_twips) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _set_run_font(
    run,
    *,
    size: float,
    cn: str = CN_BODY,
    latin: str = LATIN_BODY,
    bold: bool | None = None,
    color: str = INK,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:cs"), latin)


def _style_font(style, *, size: float, cn: str, latin: str, bold: bool, color: str) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:cs"), latin)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _style_font(normal, size=12, cn=CN_BODY, latin=LATIN_BODY, bold=False, color=INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, WD_ALIGN_PARAGRAPH.CENTER, 18, 12),
        "Heading 2": (14, WD_ALIGN_PARAGRAPH.LEFT, 14, 7),
        "Heading 3": (12, WD_ALIGN_PARAGRAPH.LEFT, 10, 5),
    }
    for name, (size, alignment, before, after) in heading_specs.items():
        style = styles[name]
        _style_font(style, size=size, cn=CN_HEADING, latin=LATIN_BODY, bold=True, color=INK)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.left_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    custom_styles = {
        "CMS TOC Title": (16, CN_HEADING, LATIN_BODY, True, INK),
        "CMS Caption": (10.5, CN_BODY, LATIN_BODY, False, MUTED),
        "CMS Table": (10.5, CN_BODY, LATIN_BODY, False, INK),
        "CMS Note": (11, CN_BODY, LATIN_BODY, False, WARNING),
        "CMS Equation": (12, CN_BODY, MATH_FONT, False, INK),
        "CMS Lead": (12, CN_BODY, LATIN_BODY, False, INK),
        "CMS Bullet": (12, CN_BODY, LATIN_BODY, False, INK),
        "CMS Number": (12, CN_BODY, LATIN_BODY, False, INK),
    }
    for name, (size, cn, latin, bold, color) in custom_styles.items():
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        _style_font(style, size=size, cn=cn, latin=latin, bold=bold, color=color)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.widow_control = True

    styles["CMS TOC Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["CMS TOC Title"].paragraph_format.space_after = Pt(14)
    styles["CMS Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["CMS Caption"].paragraph_format.keep_together = True
    styles["CMS Table"].paragraph_format.line_spacing = 1.2
    styles["CMS Table"].paragraph_format.space_after = Pt(0)
    styles["CMS Note"].paragraph_format.left_indent = Mm(4)
    styles["CMS Note"].paragraph_format.right_indent = Mm(4)
    styles["CMS Note"].paragraph_format.space_before = Pt(5)
    styles["CMS Note"].paragraph_format.space_after = Pt(8)
    styles["CMS Equation"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["CMS Equation"].paragraph_format.keep_together = True
    styles["CMS Lead"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["CMS Lead"].paragraph_format.first_line_indent = Pt(0)
    styles["CMS Lead"].paragraph_format.space_after = Pt(10)
    for list_style_name in ("CMS Bullet", "CMS Number"):
        style = styles[list_style_name]
        style.paragraph_format.left_indent = Mm(8)
        style.paragraph_format.first_line_indent = Mm(-4)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(3)


def _next_id(elements: Iterable, attribute: str) -> int:
    values = []
    for element in elements:
        value = element.get(qn(attribute))
        if value is not None and value.isdigit():
            values.append(int(value))
    return max(values, default=0) + 1


def _insert_abstract_numbering(numbering, abstract) -> None:
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
        return
    numbering.insert(list(numbering).index(first_num), abstract)


def _append_heading_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_id = max(100, _next_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId"))
    num_id = max(100, _next_id(numbering.findall(qn("w:num")), "w:numId"))

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    level_specs = (
        (0, "Heading 1", "chineseCountingThousand", "第%1章", "center", None),
        (1, "Heading 2", "decimal", "%1.%2", "left", None),
        (2, "Heading 3", "decimal", "%1.%2.%3", "left", None),
    )
    for level, style_name, number_format, text, justification, restart in level_specs:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), number_format)
        lvl.append(num_fmt)
        if restart is not None:
            restart_node = OxmlElement("w:lvlRestart")
            restart_node.set(qn("w:val"), str(restart))
            lvl.append(restart_node)
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), style_name.replace(" ", ""))
        lvl.append(p_style)
        if level > 0:
            is_legal = OxmlElement("w:isLgl")
            lvl.append(is_legal)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), justification)
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:hanging"), "0")
        p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)
    _insert_abstract_numbering(numbering, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _append_list_numbering(document: Document, *, ordered: bool, seed: int) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    abstract_id = max(seed, _next_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId"))
    num_id = max(seed, _next_id(numbering.findall(qn("w:num")), "w:numId"))
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "454")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "454")
    indent.set(qn("w:hanging"), "227")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    _insert_abstract_numbering(numbering, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if ordered:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return abstract_id, num_id


def _append_num_instance(document: Document, abstract_id: int) -> int:
    numbering = document.part.numbering_part.element
    num_id = _next_id(numbering.findall(qn("w:num")), "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def _number_heading(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def _number_list_item(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def _add_field(run, instruction: str, placeholder: str = "0") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_code = OxmlElement("w:instrText")
    field_code.set(qn("xml:space"), "preserve")
    field_code.text = f" {instruction} "
    result = OxmlElement("w:t")
    result.text = placeholder
    run._r.append(begin)
    run._r.append(field_code)
    run._r.append(separate)
    run._r.append(result)
    run._r.append(end)


def _bookmark(run, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    run._r.addprevious(start)
    run._r.addnext(end)


def _add_ref(paragraph, kind: str, identifier: str) -> None:
    run = paragraph.add_run()
    _set_run_font(run, size=12)
    bookmark = {"eq": f"eqnum_{identifier}", "fig": f"fignum_{identifier}", "table": f"tabnum_{identifier}"}[kind]
    if kind == "eq":
        paragraph.add_run("(")
    _add_field(run, f"REF {bookmark} \\h", "0")
    if kind == "eq":
        paragraph.add_run(")")


def _append_inline(paragraph, text: str) -> None:
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        reference = REF_PATTERN.fullmatch(part)
        if reference:
            _add_ref(paragraph, reference.group(1), reference.group(2))
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, size=12, bold=True)
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run, size=10.5, cn="等线", latin="Consolas", color=ACCENT)
            continue
        run = paragraph.add_run(part)
        _set_run_font(run, size=12)


def _make_anchor_from_inline(
    inline,
    *,
    width_emu: int,
    height_emu: int,
    crop_left: int = 0,
    crop_right: int = 0,
    behind_doc: bool = True,
    relative_height: int = 0,
):
    graphic = deepcopy(inline.graphic)
    if crop_left or crop_right:
        blip_fill = graphic.find(".//" + qn("pic:blipFill"))
        if blip_fill is not None:
            src_rect = OxmlElement("a:srcRect")
            src_rect.set("l", str(crop_left))
            src_rect.set("r", str(crop_right))
            stretch = blip_fill.find(qn("a:stretch"))
            if stretch is None:
                blip_fill.append(src_rect)
            else:
                blip_fill.insert(list(blip_fill).index(stretch), src_rect)

    anchor = OxmlElement("wp:anchor")
    for key, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": str(relative_height),
        "behindDoc": "1" if behind_doc else "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        anchor.set(key, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)
    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "page")
    offset_h = OxmlElement("wp:posOffset")
    offset_h.text = "0"
    position_h.append(offset_h)
    anchor.append(position_h)
    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "page")
    offset_v = OxmlElement("wp:posOffset")
    offset_v.text = "0"
    position_v.append(offset_v)
    anchor.append(position_v)
    extent = OxmlElement("wp:extent")
    extent.set("cx", str(width_emu))
    extent.set("cy", str(height_emu))
    anchor.append(extent)
    effect = OxmlElement("wp:effectExtent")
    for edge in ("l", "t", "r", "b"):
        effect.set(edge, "0")
    anchor.append(effect)
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(deepcopy(inline.docPr))
    c_nv = inline.find(qn("wp:cNvGraphicFramePr"))
    if c_nv is not None:
        anchor.append(deepcopy(c_nv))
    anchor.append(graphic)
    return anchor


def _draw_centered_segments(
    draw: ImageDraw.ImageDraw,
    *,
    y: int,
    segments: list[tuple[str, ImageFont.FreeTypeFont, tuple[int, int, int, int], int]],
) -> None:
    widths: list[int] = []
    for text, font, _fill, _offset_y in segments:
        left, _top, right, _bottom = draw.textbbox((0, 0), text, font=font)
        widths.append(right - left)
    x = (COVER_WIDTH_PX - sum(widths)) // 2
    for (text, font, fill, offset_y), width in zip(segments, widths):
        draw.text((x, y + offset_y), text, font=font, fill=fill)
        x += width


def _build_cover_image(background_path: Path, output_path: Path) -> None:
    if not SIMHEI_FONT.is_file() or not TIMES_BOLD_FONT.is_file():
        missing = [str(path) for path in (SIMHEI_FONT, TIMES_BOLD_FONT) if not path.is_file()]
        raise FileNotFoundError(f"Cover fonts not found: {', '.join(missing)}")

    with Image.open(background_path) as source:
        background = ImageOps.fit(
            source.convert("RGB"),
            (COVER_WIDTH_PX, COVER_HEIGHT_PX),
            method=Image.Resampling.LANCZOS,
            centering=(0.5, 0.5),
        ).convert("RGBA")

    background = Image.alpha_composite(
        background,
        Image.new("RGBA", background.size, (0, 0, 0, 72)),
    )
    draw = ImageDraw.Draw(background)
    title_font = ImageFont.truetype(str(TIMES_BOLD_FONT), 92)
    subtitle_latin = ImageFont.truetype(str(TIMES_BOLD_FONT), 92)
    subtitle_cn = ImageFont.truetype(str(SIMHEI_FONT), 92)
    studio_font = ImageFont.truetype(str(TIMES_BOLD_FONT), 58)
    studio_plus_font = ImageFont.truetype(str(TIMES_BOLD_FONT), 34)
    date_latin = ImageFont.truetype(str(TIMES_BOLD_FONT), 50)
    date_cn = ImageFont.truetype(str(SIMHEI_FONT), 50)

    white = (245, 249, 252, 255)
    soft_white = (220, 231, 237, 255)
    cyan = (105, 213, 226, 255)

    _draw_centered_segments(
        draw,
        y=1270,
        segments=[("Computational Mechanics Solver", title_font, white, 0)],
    )
    _draw_centered_segments(
        draw,
        y=1465,
        segments=[
            ("v1.3.2 ", subtitle_latin, cyan, 0),
            ("用户与技术说明书", subtitle_cn, cyan, 0),
        ],
    )
    _draw_centered_segments(
        draw,
        y=1890,
        segments=[
            ("Leo Li", studio_font, white, 0),
            ("+", studio_plus_font, white, -21),
            (" Studio", studio_font, white, 0),
        ],
    )
    _draw_centered_segments(
        draw,
        y=2025,
        segments=[
            ("2026 ", date_latin, soft_white, 0),
            ("年", date_cn, soft_white, 0),
            (" 7 ", date_latin, soft_white, 0),
            ("月", date_cn, soft_white, 0),
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    background.convert("RGB").save(output_path, format="PNG", optimize=True, dpi=(300, 300))


def _add_cover_background(paragraph, image_path: Path) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=A4_WIDTH, height=A4_HEIGHT)
    inline = shape._inline
    anchor = _make_anchor_from_inline(
        inline,
        width_emu=int(A4_WIDTH),
        height_emu=int(A4_HEIGHT),
        behind_doc=False,
        relative_height=251658240,
    )
    inline.getparent().replace(inline, anchor)


def _set_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def _configure_sections(document: Document) -> None:
    cover = document.sections[0]
    cover.page_width = A4_WIDTH
    cover.page_height = A4_HEIGHT
    cover.top_margin = Mm(0)
    cover.bottom_margin = Mm(0)
    cover.left_margin = Mm(0)
    cover.right_margin = Mm(0)
    cover.header_distance = Mm(0)
    cover.footer_distance = Mm(0)


def _add_cover(document: Document) -> None:
    _build_cover_image(BACKGROUND, COVER_IMAGE)
    background = document.add_paragraph()
    _add_cover_background(background, COVER_IMAGE)

    body = document.add_section(WD_SECTION.NEW_PAGE)
    body.page_width = A4_WIDTH
    body.page_height = A4_HEIGHT
    body.top_margin = Mm(22)
    body.bottom_margin = Mm(20)
    body.left_margin = Mm(25)
    body.right_margin = Mm(25)
    body.header_distance = Mm(9)
    body.footer_distance = Mm(9)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False
    _set_page_number_start(body, 1)

    header = body.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("Computational Mechanics Solver v1.3.2")
    _set_run_font(run, size=9, cn=CN_BODY, latin=LATIN_BODY, color=MUTED)
    p_pr = header._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), LINE)
    border.append(bottom)
    p_pr.append(border)

    footer = body.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    _set_run_font(run, size=9, cn=CN_BODY, latin=LATIN_BODY, color=MUTED)
    _add_field(run, "PAGE", "1")


def _add_toc(document: Document) -> None:
    paragraph = document.add_paragraph("目录", style="CMS TOC Title")
    paragraph.paragraph_format.keep_with_next = True
    toc = document.add_paragraph()
    toc.paragraph_format.first_line_indent = Pt(0)
    toc.paragraph_format.line_spacing = 1.15
    run = toc.add_run()
    _set_run_font(run, size=11)
    _add_field(run, 'TOC \\o "1-2" \\h \\z \\u', "右键更新目录")
    toc.add_run().add_break(WD_BREAK.PAGE)


def _add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="CMS Note")
    _append_inline(paragraph, text)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), WARNING_FILL)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), WARNING)
    borders.append(left)
    p_pr.append(borders)


def _caption(document: Document, *, kind: str, identifier: str, text: str, bookmark_id: int) -> None:
    paragraph = document.add_paragraph(style="CMS Caption")
    prefix = "图" if kind == "fig" else "表"
    paragraph.add_run(f"{prefix} ")
    run = paragraph.add_run()
    _set_run_font(run, size=10.5, color=MUTED)
    sequence = "图" if kind == "fig" else "表"
    _add_field(run, f"SEQ {sequence} \\* ARABIC", "0")
    _bookmark(run, f"{kind}num_{identifier}" if kind == "fig" else f"tabnum_{identifier}", bookmark_id)
    tail = paragraph.add_run(f"  {text}")
    _set_run_font(tail, size=10.5, color=MUTED)


def _add_figure(
    document: Document,
    source_dir: Path,
    path: str,
    caption: str,
    identifier: str,
    width_mm: float | None,
    bookmark_id: int,
) -> None:
    image_path = (source_dir / path).resolve()
    if not image_path.is_file():
        raise FileNotFoundError(f"Manual image not found: {image_path}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    width = Mm(min(width_mm or 155.0, 160.0))
    shape = run.add_picture(str(image_path), width=width)
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    _caption(document, kind="fig", identifier=identifier, text=caption, bookmark_id=bookmark_id)


def _column_widths(rows: list[list[str]], total_twips: int = 9072) -> list[int]:
    columns = len(rows[0])
    weights = []
    for index in range(columns):
        longest = max(len(row[index]) if index < len(row) else 0 for row in rows)
        weights.append(max(5.0, min(36.0, longest ** 0.72 + 3.0)))
    scale = total_twips / sum(weights)
    widths = [int(weight * scale) for weight in weights]
    widths[-1] += total_twips - sum(widths)
    return widths


def _add_table(document: Document, rows: list[list[str]], title: str, identifier: str, bookmark_id: int) -> None:
    _caption(document, kind="table", identifier=identifier, text=title, bookmark_id=bookmark_id)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_geometry(table, _column_widths(rows))
    _set_table_borders(table)
    _set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if row_index == 0:
                _set_cell_shading(cell, ACCENT_LIGHT)
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles["CMS Table"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            _set_run_font(run, size=10.5, bold=row_index == 0)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_equation(document: Document, identifier: str, text: str, bookmark_id: int) -> None:
    paragraph = document.add_paragraph(style="CMS Equation")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Mm(80), WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Mm(160), WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph.add_run("\t")
    equation = paragraph.add_run(text)
    _set_run_font(equation, size=12, cn=CN_BODY, latin=MATH_FONT)
    _bookmark(equation, f"eqbody_{identifier}", bookmark_id)
    paragraph.add_run("\t(")
    number = paragraph.add_run()
    _set_run_font(number, size=11)
    _add_field(number, "SEQ 公式 \\* ARABIC", "0")
    _bookmark(number, f"eqnum_{identifier}", bookmark_id + 1)
    paragraph.add_run(")")


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if index == start + 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            index += 1
            continue
        rows.append(cells)
        index += 1
    if len(rows) < 2:
        raise ValueError(f"Malformed Markdown table near line {start + 1}")
    column_count = len(rows[0])
    if any(len(row) != column_count for row in rows):
        raise ValueError(f"Inconsistent Markdown table near line {start + 1}")
    return rows, index


def _strip_front_matter(lines: list[str]) -> list[str]:
    if not lines or lines[0].strip() != "---":
        return lines
    for index in range(1, len(lines)):
        if lines[index].strip() == "---":
            return lines[index + 1 :]
    raise ValueError("Unclosed Markdown front matter")


def _build_body(
    document: Document,
    source: Path,
    heading_num_id: int,
    bullet_num_id: int,
    ordered_abstract_id: int,
    ordered_num_id: int,
) -> None:
    lines = _strip_front_matter(source.read_text(encoding="utf-8").splitlines())
    paragraph_buffer: list[str] = []
    pending_table: tuple[str, str] | None = None
    bookmark_id = 10
    ordered_list_active = False
    ordered_list_count = 0
    active_ordered_num_id = ordered_num_id

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(line.strip() for line in paragraph_buffer).strip()
        paragraph_buffer = []
        if not text:
            return
        paragraph = document.add_paragraph()
        _append_inline(paragraph, text)

    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        is_ordered_line = bool(re.match(r"^\d+\.\s+", line))
        if not is_ordered_line:
            ordered_list_active = False
        if not line:
            flush_paragraph()
            index += 1
            continue
        if line.startswith("<!--"):
            flush_paragraph()
            while index < len(lines) and "-->" not in lines[index]:
                index += 1
            index += 1
            continue
        if line == "\\newpage":
            flush_paragraph()
            document.add_page_break()
            index += 1
            continue
        if line.startswith("#"):
            flush_paragraph()
            match = re.match(r"^(#{1,3})\s+(.+)$", line)
            if not match:
                raise ValueError(f"Unsupported heading at line {index + 1}: {line}")
            level = len(match.group(1))
            paragraph = document.add_paragraph(match.group(2), style=f"Heading {level}")
            _number_heading(paragraph, heading_num_id, level - 1)
            index += 1
            continue
        image = IMAGE_PATTERN.match(line)
        if image:
            flush_paragraph()
            identifier = image.group("id") or f"figure-{bookmark_id}"
            width = float(image.group("width")) if image.group("width") else None
            _add_figure(
                document,
                source.parent,
                image.group("path"),
                image.group("caption"),
                identifier,
                width,
                bookmark_id,
            )
            bookmark_id += 1
            index += 1
            continue
        table_directive = TABLE_DIRECTIVE.match(line)
        if table_directive:
            flush_paragraph()
            pending_table = (table_directive.group("id"), table_directive.group("title"))
            index += 1
            continue
        if line.startswith("|"):
            flush_paragraph()
            if pending_table is None:
                raise ValueError(f"Table without :::table directive at line {index + 1}")
            rows, index = _parse_table(lines, index)
            _add_table(document, rows, pending_table[1], pending_table[0], bookmark_id)
            pending_table = None
            bookmark_id += 1
            continue
        equation = EQUATION_DIRECTIVE.match(line)
        if equation:
            flush_paragraph()
            _add_equation(document, equation.group("id"), equation.group("text"), bookmark_id)
            bookmark_id += 2
            index += 1
            continue
        if line.startswith("> "):
            flush_paragraph()
            note_lines = [line[2:].strip()]
            index += 1
            while index < len(lines) and lines[index].strip().startswith("> "):
                note_lines.append(lines[index].strip()[2:].strip())
                index += 1
            _add_note(document, " ".join(note_lines))
            continue
        if re.match(r"^[-*]\s+", line):
            flush_paragraph()
            paragraph = document.add_paragraph(style="CMS Bullet")
            _number_list_item(paragraph, bullet_num_id)
            _append_inline(paragraph, re.sub(r"^[-*]\s+", "", line))
            index += 1
            continue
        if is_ordered_line:
            flush_paragraph()
            if not ordered_list_active:
                if ordered_list_count > 0:
                    active_ordered_num_id = _append_num_instance(document, ordered_abstract_id)
                ordered_list_count += 1
                ordered_list_active = True
            paragraph = document.add_paragraph(style="CMS Number")
            _number_list_item(paragraph, active_ordered_num_id)
            _append_inline(paragraph, re.sub(r"^\d+\.\s+", "", line))
            index += 1
            continue
        paragraph_buffer.append(raw)
        index += 1
    flush_paragraph()
    if pending_table is not None:
        raise ValueError("Dangling :::table directive")


def _set_document_properties(document: Document) -> None:
    properties = document.core_properties
    properties.title = "Computational Mechanics Solver v1.3.2 用户与技术说明书"
    properties.subject = "二维静力学与动力学软件用户手册"
    properties.author = "Leo Li+ Studio"
    properties.keywords = "Computational Mechanics Solver, 静力学, 动力学, 用户手册"
    properties.comments = "Public user and technical manual"

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    compatibility = settings.find(qn("w:compat"))
    if compatibility is None:
        compatibility = OxmlElement("w:compat")
        settings.append(compatibility)


def build_manual(source: Path, output: Path) -> Path:
    if not source.is_file():
        raise FileNotFoundError(source)
    if not BACKGROUND.is_file():
        raise FileNotFoundError(BACKGROUND)
    document = Document()
    _set_document_properties(document)
    _configure_styles(document)
    _configure_sections(document)
    _add_cover(document)
    _add_toc(document)
    heading_num_id = _append_heading_numbering(document)
    _, bullet_num_id = _append_list_numbering(document, ordered=False, seed=101)
    ordered_abstract_id, ordered_num_id = _append_list_numbering(document, ordered=True, seed=102)
    _build_body(
        document,
        source,
        heading_num_id,
        bullet_num_id,
        ordered_abstract_id,
        ordered_num_id,
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_DOCX)
    args = parser.parse_args()
    output = build_manual(args.source.resolve(), args.output.resolve())
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
